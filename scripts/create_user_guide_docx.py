"""Generate non-technical user Word guide."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "docs" / "AeroCert-User-Quick-Guide.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("AeroCert Intelligence — Quick Guide for Users", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "For certification, quality, and engineering staff. "
        "How to open the demo and what you should see."
    )

    doc.add_heading("What this tool does", level=1)
    doc.add_paragraph(
        "Answers certification questions using your program evidence "
        "(requirements, tests, risks, changes) and shows which documents support each answer. "
        "It does not guess when evidence is missing."
    )

    doc.add_heading("Before you start", level=1)
    doc.add_paragraph("Ask IT or your project lead to start the application.")
    for item in [
        "Web page: http://localhost:3000",
        "Demo program name: dap-100",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("No API key or password is needed for the current demo.")

    doc.add_heading("Step 1 — Open the dashboard", level=1)
    doc.add_paragraph("In your browser, go to http://localhost:3000")
    doc.add_paragraph("You should see:")
    for item in [
        "Title: AeroCert Intelligence",
        "Four summary boxes at the top",
        "A text box with a sample question",
        "Button: Ask with evidence",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Step 2 — Ask a question", level=1)
    doc.add_paragraph("Type a question in plain language, then click Ask with evidence.")
    doc.add_paragraph(
        'Example: "Which unresolved certification risks could delay flight testing?"'
    )

    doc.add_heading("What you should see", level=1)
    for item in [
        "Confidence: authoritative (supporting documents were found)",
        "A short answer in plain language",
        "Evidence citations (document names and excerpts)",
        "Demo references such as RISK-CERT-07 or SW-REQ-1042",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "If nothing relevant is found, you will see insufficient_evidence. "
        "That means the system refused to guess — which is correct behavior."
    )

    doc.add_heading("Demo questions that work", level=1)
    for q in [
        "Which unresolved certification risks could delay flight testing?",
        "Show all requirements impacted by the avionics firmware update.",
    ]:
        doc.add_paragraph(q, style="List Bullet")

    doc.add_heading("If something goes wrong", level=1)
    for item in [
        "Blank page or 404: refresh the page; ask IT to restart the web app.",
        "API error: ask IT to start the backend (port 8000).",
        "No answer: confirm program dap-100 and demo documents are loaded.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("")
    note = doc.add_paragraph("Demo program DAP-100 (fictional). Technical setup: docs/TESTING.md")
    note.runs[0].italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Created {OUT}")


if __name__ == "__main__":
    main()
